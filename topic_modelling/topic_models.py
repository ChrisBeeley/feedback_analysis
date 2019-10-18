
# %% load packages
import pandas as pd
from sklearn.feature_extraction.text import TfidfVectorizer, CountVectorizer
from sklearn.decomposition import NMF, LatentDirichletAllocation
import numpy as np
import inspect

# %%

# feedback = pd.read_csv("/home/chris/Nextcloud/feedback_analysis/for_gensim.csv", encoding = "latin-1", na_values=[])

feedback = pd.read_csv("/home/chris/Nextcloud/feedback_analysis/leicester.csv", encoding = "UTF-8", na_values=[])

# grep for 8877

feedback_subset = feedback[feedback['Improve'].str.contains("8877")]

feedback_subset

feedback["word_count"] = feedback["Improve"].apply(lambda x: len(x.split()))

feedback = feedback[feedback['word_count'] > 10]

# convert the 2nd column values to a list
documents = feedback["Improve"].tolist()

# feedback.head(n = 5)

# %% write a document with the topics in

no_top_words = 8
no_top_documents = 20

from docx import Document
from docx.shared import Inches

def display_topics(H, W, feature_names, documents, no_top_words, no_top_documents):
    for topic_idx, topic in enumerate(H):
        document.add_heading("Topic %d:" % (topic_idx), 1)
        document.add_heading(" ".join([feature_names[i]
                        for i in topic.argsort()[:-no_top_words - 1:-1]]), 2)
        top_doc_indices = np.argsort( W[:,topic_idx] )[::-1][0:no_top_documents]
        for doc_index in top_doc_indices:
            document.add_paragraph(documents[doc_index])
        document.add_page_break()

# %% five topics

no_topics = 5

document = Document()

document.add_heading('Topics', 0)

# NMF is able to use tf-idf
tfidf_vectorizer = TfidfVectorizer(max_df=0.95, min_df=2, stop_words='english')
tfidf = tfidf_vectorizer.fit_transform(documents)
tfidf_feature_names = tfidf_vectorizer.get_feature_names()

# Run NMF
nmf_model = NMF(n_components=no_topics, random_state=1, alpha=.1, l1_ratio=.5, init='nndsvd').fit(tfidf)
nmf_W = nmf_model.transform(tfidf)
nmf_H = nmf_model.components_

display_topics(nmf_H, nmf_W, tfidf_feature_names, documents, no_top_words, no_top_documents)

document.save('five_topics.docx')

# %% fifteen topics

no_topics = 15

document = Document()

# Run NMF
nmf_model = NMF(n_components=no_topics, random_state=1, alpha=.1, l1_ratio=.5, init='nndsvd').fit(tfidf)
nmf_W = nmf_model.transform(tfidf)
nmf_H = nmf_model.components_

display_topics(nmf_H, nmf_W, tfidf_feature_names, documents, no_top_words, no_top_documents)

document.save('fifteen_topics.docx')

# %% 10 topics

no_topics = 10

document = Document()

"""

# df matrix

vec = CountVectorizer().fit(documents)
bag_of_words = vec.transform(documents)
sum_words = bag_of_words.sum(axis=0)
total_words = bag_of_words.sum()
# I added a bit to divide by the total number of words
words_freq = [(word, sum_words[0, idx] / total_words)  for word, idx in     vec.vocabulary_.items()]
words_freq = sorted(words_freq, key = lambda x: x[1], reverse=True)
words_freq

this is the top 200 most common words

['and', 'the', 'to', 'very', 'was', 'my', 'have', 'with', 'of', 'for', 'good', 'it', 'me', 'is', 'in', 'you', 'staff', 'been', 'has', 'all', 'be', 'service',  'this', 'not',  'are',  'at',  'so',  'help', 'care',  'that',  'on',  'as',  'would',  'were',  'friendly',  'but',  'they',  'more',  'really', 'thank', 'time', 'had', 'could', 'from', 'excellent', 'feel', 'we', 'she', 'well', 'given', 'when', 'helped', 'advice', 'support', 'food', 'her', 'about', 'better', 'what', 'can', 'information', 'made', 'great', 'do', 'after', 'get', 'exercises', 'much', 'always', 'hospital', 'found', 'like', 'nice', 'by', 'lot', 'am', 'everything', 'only', 'if', 'caring', 'us', 'an',  'there', 'professional', 'done', 'happy', 'received', 'up', 'informative', 'people', 'useful', 'lovely', 'think', 'who', 'home', 'how', 'physio', 'name', 'out', 'felt', 'treatment', 'extremely', 'no',  'nurse', 'also', 'will', 'kind', 'your', 'being', 'patient', 'health', 'explained', 'understanding', 'which', 'child', 'needs', 'because', 'enjoyed', 'needed', 'work', 'know', 'team', 'some', 'phone', 'any', 'need', 'experience', 'them', 'way', 'course', 'gave', 'first', 'or', 'questions', 'things', 'recommend', 'our', 'here', 'group', 'too', 'back', 'appointment', 'ward', 'don', 'everyone', 'lots', 'daughter', 'session', 'other', 'see', 'supportive', 'one', 'he', 'now', 'make', 'nurses', 'easy', 'did', 'family', 'person', 'improve', 'long', 'looked', 'nothing', 'problem', 'times', 'sessions', 'understand', 'life', 'their', 'patients', 'enough', 'exercise', 'over', 'weeks', 'waiting', 'pain', 'visit', 'baby', 'son', 'improved', 'different', 'brilliant', 'provided', 'talk', 'knowledgeable', 'children', 'anything', 'myself', 'stay', 'every', 'able', 'doing', 'listened', 'best', 'pleased', 'just', 'go', 'than']

"""

# now I'm going to use a custom stop word list

# NMF is able to use tf-idf
# tfidf_vectorizer = TfidfVectorizer(max_df=0.95, min_df=2, stop_words='english')
tfidf_vectorizer = TfidfVectorizer(max_df=0.95, min_df=2, stop_words=['and', 'the', 'to', 'very', 'was', 'my', 'have', 'with', 'of', 'for', 'good', 'it', 'me', 'is', 'in', 'you', 'staff', 'been', 'has', 'all', 'be', 'service',  'this', 'not',  'are',  'at',  'so',  'that',  'on',  'as',  'would',  'were', 'but',  'they',  'more',  'really', 'had', 'could', 'from', 'feel', 'we', 'she', 'well', 'given', 'when', 'her', 'about', 'better', 'what', 'can', 'made', 'great', 'do', 'after', 'get', 'much', 'always', 'found', 'like', 'nice', 'by', 'lot', 'am', 'everything', 'only', 'if', 'us', 'an',  'there', 'done', 'happy', 'received', 'up', 'people', 'useful', 'lovely', 'who', 'how', 'name', 'out', 'felt', 'extremely', 'no',  'also', 'will', 'your', 'being', 'which', 'because', 'know', 'some', 'any', 'them', 'way', 'gave', 'first', 'or', 'things', 'our', 'here', 'group', 'too', 'back', 'don', 'everyone', 'lots', 'other', 'see', 'one', 'he', 'now', 'make', 'easy', 'did', 'nothing', 'their', 'enough', 'over', 'anything', 'myself', 'every', 'able', 'doing', 'best', 'pleased', 'just', 'go', 'than'])
tfidf = tfidf_vectorizer.fit_transform(documents)
tfidf_feature_names = tfidf_vectorizer.get_feature_names()

# Run NMF
nmf_model = NMF(n_components=no_topics, random_state=1, alpha=.1, l1_ratio=.5, init='nndsvd').fit(tfidf)
nmf_W = nmf_model.transform(tfidf)
nmf_H = nmf_model.components_

display_topics(nmf_H, nmf_W, tfidf_feature_names, documents, no_top_words, no_top_documents)

document.save('ten_topics.docx')

# %% so I selected 10 topics. See other code for actual topic model output
