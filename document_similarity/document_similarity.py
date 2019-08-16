# %% markdown

## Data and package loads

# %%
import pandas as pd
import tensorflow as tf
import tensorflow_hub as hub
import matplotlib.pyplot as plt
import numpy as np
import os
import pandas as pd
import re
import seaborn as sns
import random

from IPython.display import display, Markdown

# %%
feedback = pd.read_csv("/home/chris/Nextcloud/feedback_analysis/for_gensim.csv", encoding = "latin-1", na_values=[])
feedback.head(5)

text = feedback["Improve"][0:1000].tolist()

# %%
module_url = "https://tfhub.dev/google/universal-sentence-encoder/2" #@param ["https://tfhub.dev/google/universal-sentence-encoder/2", "https://tfhub.dev/google/universal-sentence-encoder-large/3"]

# Import the Universal Sentence Encoder's TF Hub module
embed = hub.Module(module_url)

# Reduce logging output.
tf.logging.set_verbosity(tf.logging.ERROR)

with tf.Session() as session:
  session.run([tf.global_variables_initializer(), tf.tables_initializer()])
  message_embeddings = session.run(embed(text))

# %%
# this is an array of the embeddings, each is 512 long

text_array = np.array(message_embeddings).tolist()

# this is the text-  text

# dataframe

embed_frame = pd.DataFrame(
    {'text': text,
     'embedding': text_array
    })

# %%
def return_three_similar(value):
    to_sort = embed_frame["embedding"].apply(lambda i: np.inner(value, i))
    # to_sort = (-to_sort).argsort()[1:3]
    to_sort = (-to_sort).argsort()[1:5]
    return(to_sort)

list(return_three_similar(embed_frame["embedding"][0]))

embed_frame["similarity"] = embed_frame["embedding"].apply(lambda x: list(return_three_similar(x)))

# %%

row = random.sample(range(0, embed_frame.shape[0]), 10)

from docx import Document
from docx.shared import Inches

document = Document()

document.add_heading('Document similarity', 0)

for x in row:
    document.add_heading("Feedback to be compared:", 1)
    document.add_paragraph(embed_frame.text[x])
    document.add_heading("Most similar: ", 2)
    document.add_paragraph(embed_frame.text[embed_frame.similarity[x][0]])
    document.add_heading("Second most similar: ", 2)
    document.add_paragraph(embed_frame.text[embed_frame.similarity[x][1]])
    document.add_heading("Third most similar: ", 2)
    document.add_paragraph(embed_frame.text[embed_frame.similarity[x][2]])
    document.add_heading("Fourth most similar: ", 2)
    document.add_paragraph(embed_frame.text[embed_frame.similarity[x][3]])
    document.add_page_break()

document.save('document_similarity.docx')
